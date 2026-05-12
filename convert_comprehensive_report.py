#!/usr/bin/env python3
"""
Convert Comprehensive Security Report to Word (.docx) format.
"""

import os
import sys
from pathlib import Path

# Check for available conversion libraries
try:
    import pypandoc
    HAS_PYPANDOC = True
except ImportError:
    HAS_PYPANDOC = False
    print("[INFO] pypandoc not found. Install with: pip install pypandoc")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import markdown
    HAS_DOCX_CONVERSION = True
except ImportError:
    HAS_DOCX_CONVERSION = False
    print("[INFO] python-docx not found. Install with: pip install python-docx markdown")

def convert_with_pypandoc(input_file, output_file, format_type='docx'):
    """Convert markdown to Word using pypandoc."""
    try:
        if format_type == 'docx':
            output_format = 'docx'
        else:
            raise ValueError(f"Unsupported format: {format_type}")
        
        pypandoc.convert_file(
            input_file,
            output_format,
            outputfile=output_file,
            extra_args=['--standalone', '--toc']
        )
        return True
    except Exception as e:
        print(f"Error with pypandoc conversion: {e}")
        return False

def convert_markdown_to_word_enhanced(md_file, docx_file):
    """Enhanced markdown to Word converter with better formatting."""
    try:
        # Read markdown file
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Process markdown content
        lines = md_content.split('\n')
        in_code_block = False
        code_block_lines = []
        in_table = False
        table_rows = []
        
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Skip empty lines at start
            if not stripped and i == 0:
                i += 1
                continue
            
            # Code blocks
            if stripped.startswith('```'):
                if in_code_block:
                    # End code block
                    if code_block_lines:
                        code_para = doc.add_paragraph('\n'.join(code_block_lines))
                        code_para.style = 'No Spacing'
                        for run in code_para.runs:
                            run.font.name = 'Consolas'
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    code_block_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue
            
            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue
            
            # Tables
            if '|' in line and stripped.startswith('|'):
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append([cell.strip() for cell in line.split('|')[1:-1]])
                i += 1
                continue
            else:
                if in_table and table_rows:
                    # Create table
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Light Grid Accent 1'
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            # Make header row bold
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                    table_rows = []
                    in_table = False
            
            # Headers
            if stripped.startswith('# '):
                heading = doc.add_heading(stripped[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('#### '):
                doc.add_heading(stripped[5:], level=4)
            elif stripped.startswith('##### '):
                doc.add_heading(stripped[6:], level=5)
            # Horizontal rules
            elif stripped.startswith('---'):
                para = doc.add_paragraph('_' * 50)
                para.runs[0].font.color.rgb = RGBColor(192, 192, 192)
            # Lists
            elif stripped.startswith('- ') or stripped.startswith('* '):
                list_text = stripped[2:].strip()
                # Handle emojis and formatting
                para = doc.add_paragraph(list_text, style='List Bullet')
            elif stripped.startswith('1. ') or any(stripped.startswith(f'{n}. ') for n in range(2, 100)):
                # Numbered list
                num_text = stripped.split('. ', 1)
                if len(num_text) > 1:
                    para = doc.add_paragraph(num_text[1], style='List Number')
            # Checkboxes
            elif stripped.startswith('- [ ]') or stripped.startswith('- [x]') or stripped.startswith('- [X]'):
                checkbox_text = stripped[5:].strip()
                checkbox = '☐' if '[ ]' in stripped else '☑'
                doc.add_paragraph(f'{checkbox} {checkbox_text}', style='List Bullet')
            # Bold/italic text (simple detection)
            elif '**' in stripped or '__' in stripped:
                para = doc.add_paragraph()
                # Simple bold handling
                parts = stripped.replace('**', '|||').split('|||')
                for idx, part in enumerate(parts):
                    run = para.add_run(part)
                    if idx % 2 == 1:  # Odd indices are bold
                        run.bold = True
            # Regular text
            elif stripped:
                # Handle emojis and special characters
                para = doc.add_paragraph(stripped)
            else:
                # Empty line
                doc.add_paragraph()
            
            i += 1
        
        # Save document
        doc.save(docx_file)
        return True
    except Exception as e:
        print(f"Error with Word conversion: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main conversion function."""
    base_dir = Path(__file__).parent
    md_file = base_dir / 'COMPREHENSIVE_SECURITY_REPORT.md'
    
    if not md_file.exists():
        print(f"[ERROR] File not found: {md_file}")
        print(f"       Please ensure COMPREHENSIVE_SECURITY_REPORT.md exists in the project root.")
        return
    
    print(f"\n[CONVERTING] {md_file.name}")
    print("=" * 60)
    
    # Convert to Word
    docx_file = md_file.with_suffix('.docx')
    print(f"\n[CONVERTING] Converting to Word: {docx_file.name}")
    
    success = False
    if HAS_PYPANDOC:
        print("   -> Using pypandoc (best quality)...")
        success = convert_with_pypandoc(str(md_file), str(docx_file), 'docx')
        if success:
            print(f"   [SUCCESS] ✓ Word document created: {docx_file.name}")
        else:
            print(f"   [WARNING] Pypandoc failed, trying alternative method...")
            if HAS_DOCX_CONVERSION:
                success = convert_markdown_to_word_enhanced(str(md_file), str(docx_file))
                if success and docx_file.exists():
                    print(f"   [SUCCESS] ✓ Word document created (alternative method): {docx_file.name}")
    elif HAS_DOCX_CONVERSION:
        print("   -> Using python-docx...")
        success = convert_markdown_to_word_enhanced(str(md_file), str(docx_file))
        if success and docx_file.exists():
            print(f"   [SUCCESS] ✓ Word document created: {docx_file.name}")
    else:
        print(f"   [ERROR] No conversion method available.")
        print(f"   [INFO] Install dependencies:")
        print(f"          pip install pypandoc")
        print(f"          OR")
        print(f"          pip install python-docx markdown")
        return
    
    if not success:
        print(f"   [ERROR] Word conversion failed")
        return
    
    print("\n" + "=" * 60)
    print("[SUCCESS] Conversion complete!")
    if docx_file.exists():
        print(f"         Word document: {docx_file}")
        print(f"         File size: {docx_file.stat().st_size / 1024:.2f} KB")
    print("\n[NOTE] You can now open the .docx file in Microsoft Word")

if __name__ == '__main__':
    main()

