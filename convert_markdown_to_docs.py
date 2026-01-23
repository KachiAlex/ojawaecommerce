#!/usr/bin/env python3
"""
Convert Markdown files to Word (.docx) and PDF formats.
"""

import os
import sys
from pathlib import Path

try:
    import pypandoc
    HAS_PYPANDOC = True
except ImportError:
    HAS_PYPANDOC = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import markdown
    from markdown.extensions import codehilite, tables, fenced_code
    HAS_DOCX_CONVERSION = True
except ImportError:
    HAS_DOCX_CONVERSION = False

def convert_with_pypandoc(input_file, output_file, format_type='docx'):
    """Convert markdown to Word or PDF using pypandoc."""
    try:
        if format_type == 'docx':
            output_format = 'docx'
        elif format_type == 'pdf':
            output_format = 'pdf'
        else:
            raise ValueError(f"Unsupported format: {format_type}")
        
        pypandoc.convert_file(
            input_file,
            output_format,
            outputfile=output_file,
            extra_args=['--standalone']
        )
        return True
    except Exception as e:
        print(f"Error with pypandoc conversion: {e}")
        return False

def convert_markdown_to_word_simple(md_file, docx_file):
    """Simple markdown to Word converter using python-docx."""
    try:
        # Read markdown file
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML
        html = markdown.markdown(
            md_content,
            extensions=['codehilite', 'fenced_code', 'tables']
        )
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(Path(md_file).stem.replace('_', ' ').title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parse HTML and add to document (simplified approach)
        # Split by lines and process
        lines = md_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            # Code blocks
            elif line.startswith('```'):
                continue  # Skip code block markers
            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith(tuple(f'{i}. ' for i in range(1, 100))):
                doc.add_paragraph(line[line.index('.')+2:], style='List Number')
            # Regular text
            else:
                doc.add_paragraph(line)
        
        doc.save(docx_file)
        return True
    except Exception as e:
        print(f"Error with simple conversion: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main conversion function."""
    # Files to convert
    md_files = [
        'ESCROW_INTEGRATION_SUMMARY.md',
        'BANKING_PARTNER_ESCROW_INTEGRATION.md',
        'ESCROW_INTEGRATION_DIAGRAMS.md',
        'BANK_PAYMENT_FLOW_OVERVIEW.md'
    ]
    
    base_dir = Path(__file__).parent
    
    for md_file in md_files:
        md_path = base_dir / md_file
        if not md_path.exists():
            print(f"[WARNING] File not found: {md_file}")
            continue
        
        print(f"\n[CONVERTING] {md_file}")
        
        # Convert to Word
        docx_file = md_path.with_suffix('.docx')
        print(f"   -> Converting to Word: {docx_file.name}")
        
        if HAS_PYPANDOC:
            success = convert_with_pypandoc(str(md_path), str(docx_file), 'docx')
            if success:
                print(f"   [SUCCESS] Word document created: {docx_file.name}")
            else:
                print(f"   [WARNING] Pypandoc failed, trying simple converter...")
                if HAS_DOCX_CONVERSION:
                    convert_markdown_to_word_simple(str(md_path), str(docx_file))
                    if docx_file.exists():
                        print(f"   [SUCCESS] Word document created (simple method): {docx_file.name}")
        elif HAS_DOCX_CONVERSION:
            convert_markdown_to_word_simple(str(md_path), str(docx_file))
            if docx_file.exists():
                print(f"   [SUCCESS] Word document created (simple method): {docx_file.name}")
        else:
            print(f"   [ERROR] No conversion method available. Install pypandoc or python-docx")
        
        # Convert to PDF (requires pypandoc and LaTeX or wkhtmltopdf)
        pdf_file = md_path.with_suffix('.pdf')
        print(f"   -> Converting to PDF: {pdf_file.name}")
        
        if HAS_PYPANDOC:
            try:
                # Try PDF conversion
                success = convert_with_pypandoc(str(md_path), str(pdf_file), 'pdf')
                if success:
                    print(f"   [SUCCESS] PDF document created: {pdf_file.name}")
                else:
                    print(f"   [WARNING] PDF conversion failed. You may need to:")
                    print(f"      - Install wkhtmltopdf for PDF conversion, or")
                    print(f"      - Convert Word to PDF manually using Microsoft Word")
            except Exception as e:
                print(f"   [WARNING] PDF conversion error: {e}")
                print(f"      [TIP] Open the .docx file in Microsoft Word and save as PDF")
        else:
            print(f"   [WARNING] PDF conversion requires pypandoc")
            print(f"      [TIP] Open the .docx file in Microsoft Word and save as PDF")
    
    print("\n[SUCCESS] Conversion complete!")
    print("\n[NOTE] For best PDF results, open the .docx files in Microsoft Word")
    print("       and use 'Save As' -> PDF format.")

if __name__ == '__main__':
    main()

